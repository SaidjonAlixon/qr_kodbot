#!/usr/bin/env python3
"""
Namuna Word fayllar yaratish
"""
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_sample_document():
    """Namuna hujjat yaratish"""
    doc = Document()
    
    # Sarlavha
    title = doc.add_heading('Namuna Hujjat', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Paragraf
    doc.add_paragraph('Bu namuna hujjat. QR kod qo\'shilgandan keyin faylga kirish mumkin bo\'ladi.')
    
    # Ro'yxat
    doc.add_heading('Xususiyatlar:', level=1)
    doc.add_paragraph('• QR kod avtomatik qo\'shiladi', style='List Bullet')
    doc.add_paragraph('• Fayl doimiy saqlanadi', style='List Bullet')
    doc.add_paragraph('• QR kod orqali faylga kirish mumkin', style='List Bullet')
    
    # Jadval
    doc.add_heading('Ma\'lumotlar:', level=1)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    # Jadval ma'lumotlari
    table.cell(0, 0).text = 'Xususiyat'
    table.cell(0, 1).text = 'Qiymat'
    table.cell(1, 0).text = 'Fayl turi'
    table.cell(1, 1).text = 'Word hujjat'
    table.cell(2, 0).text = 'QR kod'
    table.cell(2, 1).text = 'Avtomatik qo\'shiladi'
    
    # Xulosa
    doc.add_paragraph('Bu hujjat QR kod qo\'shilgandan keyin faylga kirish uchun ishlatiladi.')
    
    return doc

def create_contract_sample():
    """Shartnoma namunasi yaratish"""
    doc = Document()
    
    # Sarlavha
    title = doc.add_heading('XIZMAT KO\'RSATISH SHARTNOMASI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Shartnoma raqami
    doc.add_paragraph('Shartnoma raqami: № 001/2025')
    doc.add_paragraph('Sana: 20.01.2025')
    doc.add_paragraph('')
    
    # Shartnoma mazmuni
    doc.add_paragraph('Quyidagi shaxslar o\'rtasida xizmat ko\'rsatish shartnomasi tuzildi:')
    doc.add_paragraph('')
    
    doc.add_paragraph('1. XIZMAT KO\'RSATUVCHI:', style='Heading 2')
    doc.add_paragraph('   Firma: "Soliq.uz" MChJ')
    doc.add_paragraph('   Manzil: Toshkent shahar')
    doc.add_paragraph('   Telefon: +998 90 123 45 67')
    doc.add_paragraph('')
    
    doc.add_paragraph('2. XIZMAT OLUVCHI:', style='Heading 2')
    doc.add_paragraph('   F.I.O: _________________')
    doc.add_paragraph('   Pasport: _________________')
    doc.add_paragraph('   Manzil: _________________')
    doc.add_paragraph('')
    
    doc.add_paragraph('3. XIZMAT TURI:', style='Heading 2')
    doc.add_paragraph('   QR kod yaratish va fayl saqlash xizmati')
    doc.add_paragraph('')
    
    doc.add_paragraph('4. SHARTLAR:', style='Heading 2')
    doc.add_paragraph('   • Xizmat pullik')
    doc.add_paragraph('   • QR kod avtomatik yaratiladi')
    doc.add_paragraph('   • Fayl doimiy saqlanadi')
    doc.add_paragraph('')
    
    # Imzolar
    doc.add_paragraph('XIZMAT KO\'RSATUVCHI:')
    doc.add_paragraph('Imzo: _________________')
    doc.add_paragraph('')
    doc.add_paragraph('XIZMAT OLUVCHI:')
    doc.add_paragraph('Imzo: _________________')
    
    return doc

def create_report_sample():
    """Hisobot namunasi yaratish"""
    doc = Document()
    
    # Sarlavha
    title = doc.add_heading('HISOBOT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Mavzu: QR kod xizmati hisoboti')
    doc.add_paragraph('Sana: 20.01.2025')
    doc.add_paragraph('')
    
    # Hisobot mazmuni
    doc.add_heading('1. Umumiy ma\'lumot', level=1)
    doc.add_paragraph('Bu hisobot QR kod xizmati ishlatilishi haqida.')
    
    doc.add_heading('2. Xizmat ko\'rsatilgan ishlar', level=1)
    doc.add_paragraph('• Fayl yuklash')
    doc.add_paragraph('• QR kod yaratish')
    doc.add_paragraph('• Fayl saqlash')
    doc.add_paragraph('• Doimiy havola yaratish')
    
    doc.add_heading('3. Natijalar', level=1)
    doc.add_paragraph('Xizmat muvaffaqiyatli ko\'rsatildi. QR kod yaratildi va fayl saqlandi.')
    
    doc.add_heading('4. Xulosa', level=1)
    doc.add_paragraph('QR kod xizmati to\'liq ishlayapti va foydalanuvchilar uchun qulay.')
    
    return doc

def main():
    """Namuna fayllar yaratish"""
    print("Namuna fayllar yaratilmoqda...")
    
    # Papka yaratish
    samples_dir = 'sample_documents'
    os.makedirs(samples_dir, exist_ok=True)
    
    # 1. Oddiy hujjat
    doc1 = create_sample_document()
    doc1.save(f'{samples_dir}/namuna_hujjat.docx')
    print("Namuna hujjat yaratildi: namuna_hujjat.docx")
    
    # 2. Shartnoma
    doc2 = create_contract_sample()
    doc2.save(f'{samples_dir}/shartnoma_namunasi.docx')
    print("Shartnoma namunasi yaratildi: shartnoma_namunasi.docx")
    
    # 3. Hisobot
    doc3 = create_report_sample()
    doc3.save(f'{samples_dir}/hisobot_namunasi.docx')
    print("Hisobot namunasi yaratildi: hisobot_namunasi.docx")
    
    print(f"\nBarcha namuna fayllar '{samples_dir}' papkasida yaratildi!")
    print("\nBu fayllarni botga yuborib, QR kod qo'shish jarayonini sinab ko'ring!")

if __name__ == '__main__':
    main()
