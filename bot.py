import os
import uuid
import qrcode
import io
import logging
import subprocess
import fitz  # PyMuPDF
from pdf2docx import Converter
from docx import Document
from docx.shared import Inches, Pt
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, CallbackQueryHandler, ContextTypes, filters
from telegram.error import BadRequest, Conflict
from functools import wraps

# Import configuration
from config import (
    TELEGRAM_BOT_TOKEN, ADMIN_TELEGRAM_ID, MAX_FILE_SIZE,
    UPLOAD_FOLDER, QR_FOLDER, ALLOWED_EXTENSIONS, 
    RAILWAY_URL, REPLIT_URL
)

# Import database functions
from database import (
    add_or_update_user, is_user_allowed, set_user_permission,
    get_all_users, add_file_record, get_all_files, get_stats
)

logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# Create directories
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(QR_FOLDER, exist_ok=True)

# Admin ID from config
ADMIN_ID = ADMIN_TELEGRAM_ID

def is_admin(user_id: int) -> bool:
    """Check if user is admin"""
    return user_id == ADMIN_ID

def require_permission(func):
    """Decorator to check if user has permission"""
    @wraps(func)
    async def wrapper(update: Update, context: ContextTypes.DEFAULT_TYPE, *args, **kwargs):
        user = update.effective_user
        user_id = user.id
        
        # Save/update user in database
        username = user.username or "No username"
        full_name = user.full_name or "No name"
        add_or_update_user(user_id, username, full_name)
        
        # Admin always has access
        if is_admin(user_id):
            return await func(update, context, *args, **kwargs)
        
        # Check if user is allowed
        if not is_user_allowed(user_id):
            await update.effective_message.reply_text(
                "‚ùå <b>Ruxsat yo'q!</b>\n\n"
                "Botdan foydalanish uchun admin ruxsati kerak.\n"
                "Iltimos admin bilan bog'laning.",
                parse_mode='HTML'
            )
            return
        
        return await func(update, context, *args, **kwargs)
    return wrapper

def get_base_url():
    """Get the base URL for file hosting"""
    # Railway da to'g'ridan-to'g'ri URL ni qaytarish
    railway_domain = os.getenv('RAILWAY_PUBLIC_DOMAIN')
    print(f"RAILWAY_PUBLIC_DOMAIN: {railway_domain}")
    print(f"RAILWAY_URL: {RAILWAY_URL}")
    
    if railway_domain:
        url = f"https://{railway_domain}"
        print(f"Using Railway domain: {url}")
        return url
    
    if RAILWAY_URL and RAILWAY_URL != 'None' and RAILWAY_URL != 'None':
        url = f"https://{RAILWAY_URL}"
        print(f"Using Railway URL: {url}")
        return url
    
    if REPLIT_URL and REPLIT_URL != 'None':
        url = f"https://{REPLIT_URL}"
        print(f"Using Replit URL: {url}")
        return url
    
    print("Using localhost fallback")
    return "http://localhost:5000"

def create_main_keyboard():
    """Create main inline keyboard"""
    keyboard = [
        [InlineKeyboardButton("üì§ Fayl yuborish", callback_data='upload')],
        [InlineKeyboardButton("üîÑ PDF ‚Üî Word", callback_data='convert_menu')],
        [InlineKeyboardButton("üìã Word faylga QR qo'shish", callback_data='add_qr_to_word')],
        [InlineKeyboardButton("üìÑ PDF faylga QR qo'shish", callback_data='add_qr_to_pdf')]
    ]
    return InlineKeyboardMarkup(keyboard)

def create_convert_keyboard():
    """Create conversion menu keyboard"""
    keyboard = [
        [InlineKeyboardButton("üìÑ PDF ‚Üí Word", callback_data='pdf_to_word')],
        [InlineKeyboardButton("üìù Word ‚Üí PDF", callback_data='word_to_pdf')],
        [InlineKeyboardButton("‚óÄÔ∏è Orqaga", callback_data='back_to_main')]
    ]
    return InlineKeyboardMarkup(keyboard)

def create_back_keyboard():
    """Create back button keyboard"""
    return InlineKeyboardMarkup([[InlineKeyboardButton("‚óÄÔ∏è Orqaga", callback_data='back_to_main')]])

@require_permission
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle /start command"""
    user = update.effective_user
    
    # Show admin badge for admin
    title = "üî± <b>Admin Panel - Soliq.uz QR Fayl Bot</b>" if is_admin(user.id) else "üåü <b>Soliq.uz QR Fayl Bot</b>"
    
    welcome_text = (
        f"{title}\n\n"
        "Bu bot orqali siz:\n"
        "‚úÖ Fayllarni yuklab, doimiy havola olishingiz\n"
        "‚úÖ Faylga QR-kod yaratishingiz\n"
        "‚úÖ QR-kodni skaner qilib faylni ochishingiz mumkin\n\n"
    )
    
    if is_admin(user.id):
        welcome_text += "üëë Admin: /admin - Admin panelni ochish\n\n"
    
    welcome_text += "Quyidagi tugmalardan birini tanlang:"
    
    await update.message.reply_text(
        welcome_text,
        reply_markup=create_main_keyboard(),
        parse_mode='HTML'
    )

@require_permission
async def button_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle button callbacks"""
    query = update.callback_query
    
    # Admin callbacks don't need permission check
    if query.data.startswith('admin_'):
        await admin_callback(update, context)
        return
    
    await query.answer()
    
    if query.data == 'upload':
        text = (
            "üì§ <b>Fayl yuklash</b>\n\n"
            "Quyidagi formatdagi fayllarni yuboring:\n"
            "üìÑ Hujjatlar: PDF, DOCX, XLSX, TXT\n"
            "üñº Rasmlar: JPG, PNG, GIF, BMP\n"
            "üì¶ Arxivlar: ZIP, RAR, 7Z\n"
            "üìä Taqdimotlar: PPTX, PPT\n\n"
            "‚ö†Ô∏è Maksimal hajm: 20MB"
        )
        keyboard = create_back_keyboard()
    elif query.data == 'convert_menu':
        text = (
            "üîÑ <b>PDF ‚Üî Word Konvertatsiya</b>\n\n"
            "Quyidagi konvertatsiya turlaridan birini tanlang:\n\n"
            "üìÑ <b>PDF ‚Üí Word</b>\n"
            "PDF faylni DOCX formatiga o'zgartirish\n\n"
            "üìù <b>Word ‚Üí PDF</b>\n"
            "DOCX faylni PDF formatiga o'zgartirish\n\n"
            "‚ö†Ô∏è Maksimal hajm: 20MB"
        )
        keyboard = create_convert_keyboard()
    elif query.data == 'pdf_to_word':
        context.user_data['convert_mode'] = 'pdf_to_word'
        text = (
            "üìÑ <b>PDF ‚Üí Word</b>\n\n"
            "Iltimos PDF faylni yuboring.\n"
            "Fayl DOCX formatiga o'zgartiriladi.\n\n"
            "‚ö†Ô∏è Maksimal hajm: 20MB"
        )
        keyboard = create_convert_keyboard()
    elif query.data == 'word_to_pdf':
        context.user_data['convert_mode'] = 'word_to_pdf'
        text = (
            "üìù <b>Word ‚Üí PDF</b>\n\n"
            "Iltimos DOCX yoki DOC faylni yuboring.\n"
            "Fayl PDF formatiga o'zgartiriladi.\n\n"
            "‚ö†Ô∏è Maksimal hajm: 20MB"
        )
        keyboard = create_convert_keyboard()
    elif query.data == 'add_qr_to_word':
        context.user_data['convert_mode'] = 'add_qr_to_word'
        text = (
            "üìã <b>Word faylga QR kod qo'shish</b>\n\n"
            "Iltimos DOCX yoki DOC faylni yuboring.\n"
            "Fayl ichiga QR kod qo'shiladi va qaytariladi.\n\n"
            "üì± QR kodni skanerlash orqali faylga kirish mumkin!\n\n"
            "‚ö†Ô∏è Maksimal hajm: 20MB"
        )
        keyboard = InlineKeyboardMarkup([[InlineKeyboardButton("‚óÄÔ∏è Orqaga", callback_data='back_to_main')]])
    elif query.data == 'add_qr_to_pdf':
        context.user_data['convert_mode'] = 'add_qr_to_pdf'
        text = (
            "üìÑ <b>PDF faylga QR kod qo'shish</b>\n\n"
            "Iltimos PDF faylni yuboring.\n"
            "Fayl ichiga QR kod qo'shiladi va qaytariladi.\n\n"
            "üì± QR kodni skanerlash orqali faylga kirish mumkin!\n\n"
            "‚ö†Ô∏è Maksimal hajm: 20MB"
        )
        keyboard = InlineKeyboardMarkup([[InlineKeyboardButton("‚óÄÔ∏è Orqaga", callback_data='back_to_main')]])
    elif query.data == 'back_to_main':
        context.user_data['convert_mode'] = None
        text = (
            "üåü <b>Soliq.uz QR Fayl Bot</b>\n\n"
            "Quyidagi tugmalardan birini tanlang:"
        )
        keyboard = create_main_keyboard()
    elif query.data == 'about':
        text = (
            "üßæ <b>Bot haqida</b>\n\n"
            "Soliq.uz QR Fayl Bot - bu fayllaringiz uchun QR-kod yaratuvchi va "
            "PDF/Word konvertatsiya xizmati.\n\n"
            "üîÑ PDF va Word formatlarini o'zgartiring\n"
            "üîí Fayllaringiz xavfsiz saqlanadi\n"
            "‚ö° Tez va qulay xizmat\n"
            "üÜì Bepul foydalanish"
        )
        keyboard = create_back_keyboard()
    elif query.data == 'contact':
        text = (
            "üìû <b>Aloqa</b>\n\n"
            "Savollaringiz bo'lsa, biz bilan bog'laning:\n\n"
            "üìß Email: support@soliq.uz\n"
            "üåê Website: https://soliq.uz\n"
            "üì± Telegram: @soliq_support"
        )
        keyboard = create_back_keyboard()
    else:
        return
    
    try:
        await query.edit_message_text(text, parse_mode='HTML', reply_markup=keyboard)
    except BadRequest as e:
        if "message is not modified" in str(e).lower():
            logger.info("Xabar allaqachon bir xil, o'zgartirish kerak emas")
        else:
            logger.error(f"Tugma bosilishida xatolik: {e}")

async def convert_pdf_to_word(pdf_path, docx_path):
    """Convert PDF to Word using pdf2docx"""
    try:
        cv = Converter(pdf_path)
        cv.convert(docx_path)
        cv.close()
        return True
    except Exception as e:
        logger.error(f"PDF to Word konvertatsiya xatoligi: {e}")
        return False

async def convert_word_to_pdf(docx_path, pdf_path):
    """Convert Word to PDF using LibreOffice"""
    try:
        output_dir = os.path.dirname(pdf_path)
        soffice_path = subprocess.run(
            ['which', 'soffice'],
            capture_output=True,
            text=True
        ).stdout.strip() or '/nix/store/s77ki6j3if918jk373md4aajqii531rd-libreoffice-24.8.7.2-wrapped/bin/soffice'
        
        result = subprocess.run(
            [soffice_path, '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_path],
            capture_output=True,
            text=True,
            timeout=60
        )
        if result.returncode == 0:
            return True
        else:
            logger.error(f"LibreOffice xatoligi: {result.stderr}")
            return False
    except subprocess.TimeoutExpired:
        logger.error("Word to PDF konvertatsiya vaqti tugadi")
        return False
    except Exception as e:
        logger.error(f"Word to PDF konvertatsiya xatoligi: {e}")
        return False

async def add_qr_to_word_document(docx_path, qr_image_path, output_path):
    """Add QR code to Word document, replace existing QR codes if found"""
    try:
        doc = Document(docx_path)
        
        # Mavjud QR kodlarni topish va o'chirish
        qr_replaced = False
        
        # Barcha paragraflarni tekshirish
        for paragraph in doc.paragraphs:
            # Paragrafdagi barcha runlarni tekshirish
            for run in paragraph.runs:
                # Agar run da rasm bo'lsa, uni o'chirish
                if run._element.xpath('.//a:blip'):
                    # Bu rasm, uni o'chirish
                    run.clear()
                    qr_replaced = True
                    print("Mavjud rasm (QR kod) topildi va o'chirildi")
        
        # Jadval ichidagi QR kodlarni tekshirish
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run._element.xpath('.//a:blip'):
                                run.clear()
                                qr_replaced = True
                                print("Jadval ichidagi mavjud rasm (QR kod) topildi va o'chirildi")
        
        # Yangi QR kod qo'shish
        if qr_replaced:
            print("Mavjud QR kod almashtirildi")
            # Pastki o'ng burchakka qo'shish
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.add_run()
            run.add_picture(qr_image_path, width=Inches(1), height=Inches(1))
        else:
            print("Mavjud QR kod topilmadi, yangi qo'shildi")
            # Agar mavjud QR kod topilmagan bo'lsa, oddiy usulda qo'shish
            if len(doc.paragraphs) > 0:
                # Add QR to the last existing paragraph (right side)
                last_paragraph = doc.paragraphs[-1]
                # Add tab to move to right side
                last_paragraph.add_run('\t')
                run = last_paragraph.add_run()
                run.add_picture(qr_image_path, width=Inches(1), height=Inches(1))
            else:
                # If document is empty, create new paragraph
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = paragraph.add_run()
                run.add_picture(qr_image_path, width=Inches(1), height=Inches(1))
        
        # Add footer to the document (all sections)
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        for section in doc.sections:
            footer = section.footer
            footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_paragraph.text = 'DIDOX.UZ Orqali tasdiqlandi!'
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save document
        doc.save(output_path)
        return True
    except Exception as e:
        logger.error(f"Word faylga QR qo'shish xatoligi: {e}")
        import traceback
        traceback.print_exc()
        return False

async def add_qr_to_pdf_document(pdf_path, qr_image_path, output_path):
    """Add QR code to PDF document, replace existing QR codes if found"""
    try:
        # Open PDF
        pdf_document = fitz.open(pdf_path)
        
        # Mavjud QR kodlarni topish va o'chirish
        qr_replaced = False
        
        # Barcha sahifalarni tekshirish
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            
            # Sahifadagi barcha rasmlarni topish
            image_list = page.get_images()
            for img_index, img in enumerate(image_list):
                # Rasm o'lchamini tekshirish (QR kod odatda kvadrat bo'ladi)
                xref = img[0]
                pix = fitz.Pixmap(pdf_document, xref)
                if pix.width == pix.height and pix.width <= 100:  # QR kod o'lchami
                    # Bu QR kod bo'lishi mumkin, uni o'chirish
                    page.delete_image(xref)
                    qr_replaced = True
                    print(f"Sahifa {page_num + 1} da mavjud QR kod topildi va o'chirildi")
                pix = None
        
        # Get last page
        last_page = pdf_document[-1]
        page_width = last_page.rect.width
        page_height = last_page.rect.height
        
        # QR code size (1x1 inch = 72x72 points)
        qr_size = 72
        
        # Position QR at bottom right (with 10pt margin)
        qr_x = page_width - qr_size - 10
        qr_y = page_height - qr_size - 10
        
        # Insert QR code image
        qr_rect = fitz.Rect(qr_x, qr_y, qr_x + qr_size, qr_y + qr_size)
        last_page.insert_image(qr_rect, filename=qr_image_path)
        
        if qr_replaced:
            print("Mavjud QR kod almashtirildi")
        else:
            print("Mavjud QR kod topilmadi, yangi qo'shildi")
        
        # Add footer text
        footer_text = "DIDOX.UZ Orqali tasdiqlandi!"
        text_position = fitz.Point(page_width / 2, page_height - 5)
        last_page.insert_text(text_position, footer_text, fontsize=10, 
                             color=(0, 0, 0), fontname="helv")
        
        # Save PDF
        pdf_document.save(output_path)
        pdf_document.close()
        return True
    except Exception as e:
        logger.error(f"PDF faylga QR qo'shish xatoligi: {e}")
        import traceback
        traceback.print_exc()
        return False

@require_permission
async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle document uploads"""
    message = update.message
    document = message.document
    user = update.effective_user
    
    if document.file_size > MAX_FILE_SIZE:
        await message.reply_text(
            "‚ùå Xatolik: Fayl hajmi 20MB dan oshmasligi kerak!",
            reply_markup=create_back_keyboard()
        )
        return
    
    file_extension = document.file_name.split('.')[-1].lower()
    convert_mode = context.user_data.get('convert_mode')
    
    if convert_mode == 'pdf_to_word':
        if file_extension != 'pdf':
            await message.reply_text(
                "‚ùå Xatolik: Iltimos PDF fayl yuboring!",
                reply_markup=create_convert_keyboard()
            )
            return
        
        status_message = await message.reply_text("‚è≥ PDF Word ga o'zgartrilmoqda...")
        
        pdf_path = None
        docx_path = None
        try:
            file = await context.bot.get_file(document.file_id)
            unique_id = str(uuid.uuid4())
            pdf_path = os.path.join(UPLOAD_FOLDER, f"{unique_id}.pdf")
            docx_path = os.path.join(UPLOAD_FOLDER, f"{unique_id}.docx")
            
            await file.download_to_drive(pdf_path)
            
            success = await convert_pdf_to_word(pdf_path, docx_path)
            
            if success and os.path.exists(docx_path):
                await status_message.edit_text("‚úÖ Konvertatsiya muvaffaqiyatli!")
                
                # Create URL and save to database
                docx_filename = f"{unique_id}.docx"
                file_url = f"{get_base_url()}/files/{docx_filename}"
                file_size = os.path.getsize(docx_path)
                
                try:
                    add_file_record(
                        user_id=user.id,
                        file_name=f"{os.path.splitext(document.file_name)[0]}.docx",
                        file_path=docx_path,
                        file_url=file_url,
                        file_type='docx',
                        file_size=file_size,
                        service_used='pdf_to_word'
                    )
                    logger.info(f"PDF to Word conversion saved: {document.file_name} by user {user.id}")
                except Exception as e:
                    logger.error(f"Failed to save PDF to Word record: {e}")
                
                with open(docx_path, 'rb') as docx_file:
                    await message.reply_document(
                        document=docx_file,
                        filename=f"{os.path.splitext(document.file_name)[0]}.docx",
                        caption="‚úÖ PDF Word formatiga o'zgartirildi\nüåê Soliq.uz",
                        reply_markup=create_convert_keyboard()
                    )
                context.user_data['convert_mode'] = None
            else:
                await status_message.edit_text(
                    "‚ùå Konvertatsiya xatoligi. Iltimos qaytadan urinib ko'ring.",
                    reply_markup=create_convert_keyboard()
                )
        except Exception as e:
            logger.error(f"PDF to Word handler xatoligi: {e}")
            await status_message.edit_text(
                f"‚ùå Xatolik yuz berdi: {str(e)}",
                reply_markup=create_convert_keyboard()
            )
        finally:
            # Clean up only the source PDF, keep the converted DOCX
            if pdf_path and os.path.exists(pdf_path):
                os.remove(pdf_path)
        return
    
    elif convert_mode == 'word_to_pdf':
        if file_extension not in ['docx', 'doc']:
            await message.reply_text(
                "‚ùå Xatolik: Iltimos DOCX yoki DOC fayl yuboring!",
                reply_markup=create_convert_keyboard()
            )
            return
        
        status_message = await message.reply_text("‚è≥ Word PDF ga o'zgartrilmoqda...")
        
        docx_path = None
        pdf_path = None
        try:
            file = await context.bot.get_file(document.file_id)
            unique_id = str(uuid.uuid4())
            docx_path = os.path.join(UPLOAD_FOLDER, f"{unique_id}.{file_extension}")
            pdf_filename = f"{unique_id}.pdf"
            pdf_path = os.path.join(UPLOAD_FOLDER, pdf_filename)
            
            await file.download_to_drive(docx_path)
            
            success = await convert_word_to_pdf(docx_path, pdf_path)
            
            if success and os.path.exists(pdf_path):
                await status_message.edit_text("‚úÖ Konvertatsiya muvaffaqiyatli!")
                
                # Create URL and save to database
                file_url = f"{get_base_url()}/files/{pdf_filename}"
                file_size = os.path.getsize(pdf_path)
                
                try:
                    add_file_record(
                        user_id=user.id,
                        file_name=f"{os.path.splitext(document.file_name)[0]}.pdf",
                        file_path=pdf_path,
                        file_url=file_url,
                        file_type='pdf',
                        file_size=file_size,
                        service_used='word_to_pdf'
                    )
                    logger.info(f"Word to PDF conversion saved: {document.file_name} by user {user.id}")
                except Exception as e:
                    logger.error(f"Failed to save Word to PDF record: {e}")
                
                with open(pdf_path, 'rb') as pdf_file:
                    await message.reply_document(
                        document=pdf_file,
                        filename=f"{os.path.splitext(document.file_name)[0]}.pdf",
                        caption="‚úÖ Word PDF formatiga o'zgartirildi\nüåê Soliq.uz",
                        reply_markup=create_convert_keyboard()
                    )
                context.user_data['convert_mode'] = None
            else:
                await status_message.edit_text(
                    "‚ùå Konvertatsiya xatoligi. Iltimos qaytadan urinib ko'ring.",
                    reply_markup=create_convert_keyboard()
                )
        except Exception as e:
            logger.error(f"Word to PDF handler xatoligi: {e}")
            await status_message.edit_text(
                f"‚ùå Xatolik yuz berdi: {str(e)}",
                reply_markup=create_convert_keyboard()
            )
        finally:
            # Clean up only the source DOCX, keep the converted PDF
            if docx_path and os.path.exists(docx_path):
                os.remove(docx_path)
        return
    
    elif convert_mode == 'add_qr_to_word':
        if file_extension not in ['docx', 'doc']:
            await message.reply_text(
                "‚ùå Xatolik: Iltimos DOCX yoki DOC fayl yuboring!",
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("‚óÄÔ∏è Orqaga", callback_data='back_to_main')]])
            )
            return
        
        status_message = await message.reply_text("‚è≥ Word faylga QR kod qo'shilmoqda...")
        
        original_file_path = None
        converted_docx_path = None
        qr_image_path = None
        output_docx_path = None
        try:
            file = await context.bot.get_file(document.file_id)
            unique_id = str(uuid.uuid4())
            original_file_path = os.path.join(UPLOAD_FOLDER, f"{unique_id}_original.{file_extension}")
            output_docx_path = os.path.join(UPLOAD_FOLDER, f"{unique_id}_with_qr.docx")
            qr_image_path = os.path.join(QR_FOLDER, f"{unique_id}.png")
            
            # Download original file
            await file.download_to_drive(original_file_path)
            
            # If DOC, convert to DOCX first
            if file_extension == 'doc':
                await status_message.edit_text("‚è≥ DOC faylni DOCX ga o'zgartirish...")
                converted_docx_path = os.path.join(UPLOAD_FOLDER, f"{unique_id}_converted.docx")
                
                soffice_path = subprocess.run(
                    ['which', 'soffice'],
                    capture_output=True,
                    text=True
                ).stdout.strip() or '/nix/store/s77ki6j3if918jk373md4aajqii531rd-libreoffice-24.8.7.2-wrapped/bin/soffice'
                
                result = subprocess.run(
                    [soffice_path, '--headless', '--convert-to', 'docx', '--outdir', UPLOAD_FOLDER, original_file_path],
                    capture_output=True,
                    text=True,
                    timeout=60
                )
                
                if result.returncode != 0:
                    await status_message.edit_text(
                        "‚ùå DOC ni DOCX ga konvertatsiya qilishda xatolik.",
                        reply_markup=create_back_keyboard()
                    )
                    return
                
                # LibreOffice creates file with same base name but .docx extension
                converted_docx_path = os.path.join(UPLOAD_FOLDER, f"{unique_id}_original.docx")
                working_docx_path = converted_docx_path
                await status_message.edit_text("‚è≥ QR kod qo'shilmoqda...")
            else:
                working_docx_path = original_file_path
            
            # Create permanent file link and QR code
            permanent_filename = f"{uuid.uuid4()}.docx"
            permanent_file_path = os.path.join(UPLOAD_FOLDER, permanent_filename)
            file_url = f"{get_base_url()}/files/{permanent_filename}"
            
            # Generate QR code
            qr = qrcode.QRCode(
                version=1,
                error_correction=qrcode.constants.ERROR_CORRECT_L,
                box_size=10,
                border=4,
            )
            qr.add_data(file_url)
            qr.make(fit=True)
            
            img = qr.make_image(fill_color="black", back_color="white")
            img.save(qr_image_path)
            
            # Add QR code to Word document
            success = await add_qr_to_word_document(working_docx_path, qr_image_path, output_docx_path)
            
            if success and os.path.exists(output_docx_path):
                await status_message.edit_text("‚úÖ QR kod muvaffaqiyatli qo'shildi!")
                
                # Save the file with QR code as the permanent file
                os.rename(output_docx_path, permanent_file_path)
                
                # Save to database
                file_size = os.path.getsize(permanent_file_path)
                try:
                    add_file_record(
                        user_id=user.id,
                        file_name=f"{os.path.splitext(document.file_name)[0]}_QR.docx",
                        file_path=permanent_file_path,
                        file_url=file_url,
                        file_type='docx',
                        file_size=file_size,
                        service_used='qr_to_word'
                    )
                    logger.info(f"QR to Word saved: {document.file_name} by user {user.id}")
                except Exception as e:
                    logger.error(f"Failed to save QR to Word record: {e}")
                
                # Send document with QR code
                caption_text = "‚úÖ Word faylga QR kod qo'shildi!\n\n"
                if qr_replaced:
                    caption_text += "üîÑ Mavjud QR kod almashtirildi!\n\n"
                else:
                    caption_text += "‚ûï Yangi QR kod qo'shildi!\n\n"
                caption_text += f"üì• Yuklab olish: {file_url}\nüåê Soliq.uz"
                
                with open(permanent_file_path, 'rb') as docx_file:
                    await message.reply_document(
                        document=docx_file,
                        filename=f"{os.path.splitext(document.file_name)[0]}_QR.docx",
                        caption=caption_text,
                        reply_markup=create_back_keyboard()
                    )
                context.user_data['convert_mode'] = None
            else:
                await status_message.edit_text(
                    "‚ùå QR kod qo'shishda xatolik. Iltimos qaytadan urinib ko'ring.",
                    reply_markup=create_back_keyboard()
                )
        except Exception as e:
            logger.error(f"Word faylga QR qo'shish handler xatoligi: {e}")
            await status_message.edit_text(
                f"‚ùå Xatolik yuz berdi: {str(e)}",
                reply_markup=create_back_keyboard()
            )
        finally:
            if original_file_path and os.path.exists(original_file_path):
                os.remove(original_file_path)
            if converted_docx_path and os.path.exists(converted_docx_path):
                os.remove(converted_docx_path)
            if qr_image_path and os.path.exists(qr_image_path):
                os.remove(qr_image_path)
            if output_docx_path and os.path.exists(output_docx_path):
                os.remove(output_docx_path)
        return
    
    elif convert_mode == 'add_qr_to_pdf':
        if file_extension != 'pdf':
            await message.reply_text(
                "‚ùå Xatolik: Iltimos PDF fayl yuboring!",
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("‚óÄÔ∏è Orqaga", callback_data='back_to_main')]])
            )
            return
        
        status_message = await message.reply_text("‚è≥ PDF faylga QR kod qo'shilmoqda...")
        
        original_pdf_path = None
        qr_image_path = None
        output_pdf_path = None
        try:
            file = await context.bot.get_file(document.file_id)
            unique_id = str(uuid.uuid4())
            original_pdf_path = os.path.join(UPLOAD_FOLDER, f"{unique_id}_original.pdf")
            output_pdf_path = os.path.join(UPLOAD_FOLDER, f"{unique_id}_with_qr.pdf")
            qr_image_path = os.path.join(QR_FOLDER, f"{unique_id}.png")
            
            # Download original file
            await file.download_to_drive(original_pdf_path)
            
            # Create permanent file link and QR code
            permanent_filename = f"{uuid.uuid4()}.pdf"
            permanent_file_path = os.path.join(UPLOAD_FOLDER, permanent_filename)
            file_url = f"{get_base_url()}/files/{permanent_filename}"
            
            # Generate QR code
            qr = qrcode.QRCode(
                version=1,
                error_correction=qrcode.constants.ERROR_CORRECT_L,
                box_size=10,
                border=4,
            )
            qr.add_data(file_url)
            qr.make(fit=True)
            
            img = qr.make_image(fill_color="black", back_color="white")
            img.save(qr_image_path)
            
            # Add QR code to PDF document
            success = await add_qr_to_pdf_document(original_pdf_path, qr_image_path, output_pdf_path)
            
            if success and os.path.exists(output_pdf_path):
                await status_message.edit_text("‚úÖ QR kod muvaffaqiyatli qo'shildi!")
                
                # Save the file with QR code as the permanent file
                os.rename(output_pdf_path, permanent_file_path)
                
                # Save to database
                file_size = os.path.getsize(permanent_file_path)
                try:
                    add_file_record(
                        user_id=user.id,
                        file_name=f"{os.path.splitext(document.file_name)[0]}_QR.pdf",
                        file_path=permanent_file_path,
                        file_url=file_url,
                        file_type='pdf',
                        file_size=file_size,
                        service_used='qr_to_pdf'
                    )
                    logger.info(f"QR to PDF saved: {document.file_name} by user {user.id}")
                except Exception as e:
                    logger.error(f"Failed to save QR to PDF record: {e}")
                
                # Send document with QR code
                caption_text = "‚úÖ PDF faylga QR kod qo'shildi!\n\n"
                if qr_replaced:
                    caption_text += "üîÑ Mavjud QR kod almashtirildi!\n\n"
                else:
                    caption_text += "‚ûï Yangi QR kod qo'shildi!\n\n"
                caption_text += f"üì• Yuklab olish: {file_url}\nüåê Soliq.uz"
                
                with open(permanent_file_path, 'rb') as pdf_file:
                    await message.reply_document(
                        document=pdf_file,
                        filename=f"{os.path.splitext(document.file_name)[0]}_QR.pdf",
                        caption=caption_text,
                        reply_markup=create_back_keyboard()
                    )
                context.user_data['convert_mode'] = None
            else:
                await status_message.edit_text(
                    "‚ùå QR kod qo'shishda xatolik. Iltimos qaytadan urinib ko'ring.",
                    reply_markup=create_back_keyboard()
                )
        except Exception as e:
            logger.error(f"PDF faylga QR qo'shish handler xatoligi: {e}")
            await status_message.edit_text(
                f"‚ùå Xatolik yuz berdi: {str(e)}",
                reply_markup=create_back_keyboard()
            )
        finally:
            if original_pdf_path and os.path.exists(original_pdf_path):
                os.remove(original_pdf_path)
            if qr_image_path and os.path.exists(qr_image_path):
                os.remove(qr_image_path)
            if output_pdf_path and os.path.exists(output_pdf_path):
                os.remove(output_pdf_path)
        return
    
    if file_extension not in ALLOWED_EXTENSIONS:
        await message.reply_text(
            f"‚ùå Xatolik: '{file_extension}' formatidagi fayllar qo'llab-quvvatlanmaydi!",
            reply_markup=create_back_keyboard()
        )
        return
    
    status_message = await message.reply_text("‚è≥ Fayl yuklanmoqda...")
    
    try:
        file = await context.bot.get_file(document.file_id)
        unique_filename = f"{uuid.uuid4()}.{file_extension}"
        file_path = os.path.join(UPLOAD_FOLDER, unique_filename)
        
        await file.download_to_drive(file_path)
        
        file_url = f"{get_base_url()}/files/{unique_filename}"
        
        # Save file record to database
        try:
            add_file_record(
                user_id=user.id,
                file_name=document.file_name,
                file_path=file_path,
                file_url=file_url,
                file_type=file_extension,
                file_size=document.file_size
            )
            logger.info(f"File record saved: {document.file_name} by user {user.id}")
        except Exception as e:
            logger.error(f"Failed to save file record: {e}")
        
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=4,
        )
        qr.add_data(file_url)
        qr.make(fit=True)
        
        img = qr.make_image(fill_color="black", back_color="white")
        
        img_byte_arr = io.BytesIO()
        img.save(img_byte_arr, format='PNG')
        img_byte_arr.seek(0)
        
        await status_message.edit_text("‚úÖ Fayl muvaffaqiyatly yuklandi!")
        
        success_text = (
            f"‚úÖ <b>Faylingiz muvaffaqiyatli yuklandi!</b>\n\n"
            f"üìÑ Fayl nomi: {document.file_name}\n"
            f"üìä Hajmi: {document.file_size / 1024:.2f} KB\n\n"
            f"üîó <b>Faylga havola:</b>\n{file_url}\n\n"
            f"üìé QR-kodni skaner qiling yoki havolani bosing:"
        )
        
        await message.reply_text(success_text, parse_mode='HTML')
        
        await message.reply_photo(
            photo=img_byte_arr,
            caption=f"üì± QR-kodni skaner qilish orqali faylni oching\nüåê Soliq.uz",
            reply_markup=create_back_keyboard()
        )
        
    except Exception as e:
        await status_message.edit_text(
            f"‚ùå Xatolik yuz berdi: {str(e)}",
            reply_markup=create_back_keyboard()
        )

@require_permission
async def handle_photo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle photo uploads"""
    message = update.message
    photo = message.photo[-1]
    user = update.effective_user
    
    if photo.file_size > MAX_FILE_SIZE:
        await message.reply_text(
            "‚ùå Xatolik: Rasm hajmi 20MB dan oshmasligi kerak!",
            reply_markup=create_back_keyboard()
        )
        return
    
    status_message = await message.reply_text("‚è≥ Rasm yuklanmoqda...")
    
    try:
        file = await context.bot.get_file(photo.file_id)
        unique_filename = f"{uuid.uuid4()}.jpg"
        file_path = os.path.join(UPLOAD_FOLDER, unique_filename)
        
        await file.download_to_drive(file_path)
        
        file_url = f"{get_base_url()}/files/{unique_filename}"
        
        # Save file record to database
        try:
            add_file_record(
                user_id=user.id,
                file_name=f"photo_{unique_filename}",
                file_path=file_path,
                file_url=file_url,
                file_type='jpg',
                file_size=photo.file_size
            )
            logger.info(f"Photo record saved: photo_{unique_filename} by user {user.id}")
        except Exception as e:
            logger.error(f"Failed to save photo record: {e}")
        
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=4,
        )
        qr.add_data(file_url)
        qr.make(fit=True)
        
        img = qr.make_image(fill_color="black", back_color="white")
        
        img_byte_arr = io.BytesIO()
        img.save(img_byte_arr, format='PNG')
        img_byte_arr.seek(0)
        
        await status_message.edit_text("‚úÖ Rasm muvaffaqiyatly yuklandi!")
        
        success_text = (
            f"‚úÖ <b>Rasmingiz muvaffaqiyatli yuklandi!</b>\n\n"
            f"üìä Hajmi: {photo.file_size / 1024:.2f} KB\n\n"
            f"üîó <b>Rasmga havola:</b>\n{file_url}\n\n"
            f"üìé QR-kodni skaner qiling yoki havolani bosing:"
        )
        
        await message.reply_text(success_text, parse_mode='HTML')
        
        await message.reply_photo(
            photo=img_byte_arr,
            caption=f"üì± QR-kodni skaner qilish orqali rasmni oching\nüåê Soliq.uz",
            reply_markup=create_back_keyboard()
        )
        
    except Exception as e:
        await status_message.edit_text(
            f"‚ùå Xatolik yuz berdi: {str(e)}",
            reply_markup=create_back_keyboard()
        )

async def error_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle errors in the bot"""
    logger.error(f"Xatolik yuz berdi: {context.error}")
    
    if isinstance(context.error, Conflict):
        logger.warning("Conflict xatoligi: Bir nechta bot nusxasi ishlayotgan bo'lishi mumkin")
        return
    
    try:
        if update and update.effective_message:
            await update.effective_message.reply_text(
                "‚ùå Uzr, xatolik yuz berdi. Iltimos qaytadan urinib ko'ring.",
                reply_markup=create_back_keyboard()
            )
    except Exception as e:
        logger.error(f"Xatolik xabarini yuborishda muammo: {e}")

async def admin_panel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Admin panel - only for admin"""
    user_id = update.effective_user.id
    
    if not is_admin(user_id):
        await update.message.reply_text("‚ùå Bu buyruq faqat admin uchun!")
        return
    
    stats = get_stats()
    
    text = (
        "üî± <b>ADMIN PANEL</b>\n\n"
        f"üë• Jami foydalanuvchilar: {stats['total_users']}\n"
        f"‚úÖ Ruxsat berilganlar: {stats['allowed_users']}\n"
        f"üìÅ Jami fayllar: {stats['total_files']}\n"
        f"üíæ Jami hajm: {stats['total_size'] / (1024*1024):.2f} MB\n\n"
        "Quyidagi amallardan birini tanlang:"
    )
    
    keyboard = [
        [InlineKeyboardButton("üë• Foydalanuvchilar", callback_data='admin_users')],
        [InlineKeyboardButton("üìÇ Yuklangan fayllar", callback_data='admin_files')],
        [InlineKeyboardButton("‚óÄÔ∏è Orqaga", callback_data='admin_close')]
    ]
    
    await update.message.reply_text(
        text,
        reply_markup=InlineKeyboardMarkup(keyboard),
        parse_mode='HTML'
    )

async def admin_users_list(query, context):
    """Show users list for admin"""
    users = get_all_users()
    
    if not users:
        await query.edit_message_text(
            "üë• Foydalanuvchilar ro'yxati bo'sh",
            reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("‚óÄÔ∏è Orqaga", callback_data='admin_back')]])
        )
        return
    
    text = "üë• <b>Foydalanuvchilar ro'yxati:</b>\n\n"
    keyboard = []
    
    for user in users[:20]:  # Show first 20 users
        user_id_db, username, full_name, is_allowed, created_at = user
        status = "‚úÖ" if is_allowed else "‚ùå"
        text += f"{status} <code>{user_id_db}</code> - {full_name} (@{username})\n"
        
        # Add button for each user
        button_text = f"{'üîì' if is_allowed else 'üîí'} {full_name[:15]}"
        keyboard.append([InlineKeyboardButton(button_text, callback_data=f'admin_toggle_{user_id_db}')])
    
    keyboard.append([InlineKeyboardButton("‚óÄÔ∏è Orqaga", callback_data='admin_back')])
    
    await query.edit_message_text(
        text,
        reply_markup=InlineKeyboardMarkup(keyboard),
        parse_mode='HTML'
    )

async def admin_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle admin panel callbacks"""
    query = update.callback_query
    await query.answer()
    
    user_id = query.from_user.id
    
    if not is_admin(user_id):
        await query.edit_message_text("‚ùå Ruxsat yo'q!")
        return
    
    if query.data == 'admin_users':
        await admin_users_list(query, context)
    
    elif query.data == 'admin_files':
        try:
            files = get_all_files()
            logger.info(f"Admin panel - files count: {len(files)}")
            
            if not files:
                await query.edit_message_text(
                    "üìÇ <b>Fayllar ro'yxati bo'sh</b>\n\n"
                    "‚ÑπÔ∏è Hali hech kim fayl yuklamagan.\n"
                    "Foydalanuvchilar fayl yuk lagach, bu yerda ko'rsatiladi.",
                    reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("‚óÄÔ∏è Orqaga", callback_data='admin_back')]]),
                    parse_mode='HTML'
                )
                return
            
            text = "üìÇ <b>Yuklangan fayllar:</b>\n\n"
            
            # Map service names to Uzbek
            service_names = {
                'file_upload': 'üì§ Fayl yuklash',
                'pdf_to_word': 'üìÑ PDF ‚Üí Word',
                'word_to_pdf': 'üìÑ Word ‚Üí PDF',
                'qr_to_word': 'üî≤ QR ‚Üí Word',
                'qr_to_pdf': 'üî≤ QR ‚Üí PDF'
            }
            
            for file in files[:15]:  # Show first 15 files
                file_id, file_name, file_url, file_type, file_size, service_used, uploaded_at, username, full_name = file
                size_mb = file_size / (1024 * 1024)
                service_name = service_names.get(service_used, service_used)
                text += f"üìÑ <b>{file_name}</b>\n"
                text += f"üë§ {full_name} (@{username})\n"
                text += f"üîß Xizmat: {service_name}\n"
                text += f"üìä {size_mb:.2f} MB | {file_type.upper()}\n"
                text += f"üîó {file_url}\n"
                text += f"üìÖ {uploaded_at}\n\n"
            
            keyboard = [[InlineKeyboardButton("‚óÄÔ∏è Orqaga", callback_data='admin_back')]]
            
            await query.edit_message_text(
                text,
                reply_markup=InlineKeyboardMarkup(keyboard),
                parse_mode='HTML'
            )
        except Exception as e:
            logger.error(f"Admin files panel error: {e}")
            await query.edit_message_text(
                f"‚ùå Xatolik yuz berdi: {str(e)}",
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("‚óÄÔ∏è Orqaga", callback_data='admin_back')]])
            )
    
    elif query.data.startswith('admin_toggle_'):
        # Show user detail page
        target_user_id = int(query.data.split('_')[2])
        
        # Get user info from database
        users = get_all_users()
        user_info = None
        for user in users:
            if user[0] == target_user_id:
                user_info = user
                break
        
        if not user_info:
            await query.answer("‚ùå Foydalanuvchi topilmadi!")
            return
        
        user_id_db, username, full_name, is_allowed, created_at = user_info
        status = "‚úÖ Ruxsat berilgan" if is_allowed else "‚ùå Ruxsat yo'q"
        
        text = (
            f"üë§ <b>Foydalanuvchi ma'lumotlari:</b>\n\n"
            f"üÜî ID: <code>{user_id_db}</code>\n"
            f"üë®‚Äçüíº Ism: {full_name}\n"
            f"üì± Username: @{username}\n"
            f"üìä Holat: {status}\n"
            f"üìÖ Qo'shildi: {created_at}\n\n"
            f"Foydalanuvchiga ruxsat bering yoki rad eting:"
        )
        
        keyboard = [
            [InlineKeyboardButton("‚úÖ Ruxsat berish", callback_data=f'admin_grant_{user_id_db}')],
            [InlineKeyboardButton("‚ùå Rad etish", callback_data=f'admin_deny_{user_id_db}')],
            [InlineKeyboardButton("‚óÄÔ∏è Orqaga", callback_data='admin_users')]
        ]
        
        await query.edit_message_text(
            text,
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='HTML'
        )
    
    elif query.data.startswith('admin_grant_'):
        # Grant permission
        target_user_id = int(query.data.split('_')[2])
        
        # Check current status
        is_allowed_now = is_user_allowed(target_user_id)
        
        if is_allowed_now:
            await query.answer("‚ÑπÔ∏è Foydalanuvchi allaqachon ruxsat berilgan!", show_alert=True)
        else:
            # Grant permission
            set_user_permission(target_user_id, True)
            
            # Send notification to user
            try:
                await context.bot.send_message(
                    chat_id=target_user_id,
                    text="‚úÖ <b>Tabriklaymiz!</b>\n\n"
                         "Sizga botdan foydalanish uchun ruxsat berildi.\n\n"
                         "Botdan foydalanish uchun /start buyrug'ini bosing.",
                    parse_mode='HTML'
                )
                await query.answer("‚úÖ Ruxsat berildi va foydalanuvchiga xabar yuborildi!", show_alert=True)
            except Exception as e:
                logger.error(f"Foydalanuvchiga xabar yuborishda xato: {e}")
                await query.answer("‚úÖ Ruxsat berildi (lekin xabar yuborilmadi)", show_alert=True)
        
        # Go back to users list
        await admin_users_list(query, context)
    
    elif query.data.startswith('admin_deny_'):
        # Deny permission
        target_user_id = int(query.data.split('_')[2])
        
        # Check current status
        is_allowed_now = is_user_allowed(target_user_id)
        
        if not is_allowed_now:
            await query.answer("‚ÑπÔ∏è Foydalanuvchi allaqachon rad etilgan!", show_alert=True)
        else:
            # Deny permission
            set_user_permission(target_user_id, False)
            
            # Send notification to user
            try:
                await context.bot.send_message(
                    chat_id=target_user_id,
                    text="‚ùå <b>Xabarnoma</b>\n\n"
                         "Sizning botdan foydalanish ruxsatingiz bekor qilindi.\n\n"
                         "Agar bu xato deb hisoblasangiz, admin bilan bog'laning.",
                    parse_mode='HTML'
                )
                await query.answer("‚ùå Ruxsat bekor qilindi va foydalanuvchiga xabar yuborildi!", show_alert=True)
            except Exception as e:
                logger.error(f"Foydalanuvchiga xabar yuborishda xato: {e}")
                await query.answer("‚ùå Ruxsat bekor qilindi (lekin xabar yuborilmadi)", show_alert=True)
        
        # Go back to users list
        await admin_users_list(query, context)
    
    elif query.data == 'admin_back':
        stats = get_stats()
        
        text = (
            "üî± <b>ADMIN PANEL</b>\n\n"
            f"üë• Jami foydalanuvchilar: {stats['total_users']}\n"
            f"‚úÖ Ruxsat berilganlar: {stats['allowed_users']}\n"
            f"üìÅ Jami fayllar: {stats['total_files']}\n"
            f"üíæ Jami hajm: {stats['total_size'] / (1024*1024):.2f} MB\n\n"
            "Quyidagi amallardan birini tanlang:"
        )
        
        keyboard = [
            [InlineKeyboardButton("üë• Foydalanuvchilar", callback_data='admin_users')],
            [InlineKeyboardButton("üìÇ Yuklangan fayllar", callback_data='admin_files')],
            [InlineKeyboardButton("‚óÄÔ∏è Yopish", callback_data='admin_close')]
        ]
        
        await query.edit_message_text(
            text,
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='HTML'
        )
    
    elif query.data == 'admin_close':
        await query.edit_message_text("‚úÖ Admin panel yopildi")

def main():
    """Main function to run the bot"""
    print(f"Bot main() funksiyasi ishga tushdi...")
    print(f"TELEGRAM_BOT_TOKEN: {TELEGRAM_BOT_TOKEN[:10] if TELEGRAM_BOT_TOKEN else 'None'}...")
    
    # File server ni alohida thread da ishga tushirish
    import threading
    import time
    
    def start_file_server():
        """File server ni ishga tushirish"""
        print("File server ishga tushmoqda...")
        try:
            # PORT ni to'g'ridan-to'g'ri olish
            port = int(os.getenv('PORT', '5000'))
            print(f"File server port: {port}")
            
            import file_server
            file_server.app.run(host='0.0.0.0', port=port, debug=False, use_reloader=False)
        except Exception as e:
            print(f"File server xatoligi: {e}")
            import traceback
            traceback.print_exc()
    
    # File server ni background da ishga tushirish
    file_server_thread = threading.Thread(target=start_file_server, daemon=True)
    file_server_thread.start()
    time.sleep(2)
    
    if not TELEGRAM_BOT_TOKEN or TELEGRAM_BOT_TOKEN == 'YOUR_BOT_TOKEN_HERE':
        print("XATOLIK: TELEGRAM_BOT_TOKEN muhit o'zgaruvchisi topilmadi!")
        print("Botni ishga tushirish uchun Telegram Bot Token kerak.")
        print("config.py faylida TELEGRAM_BOT_TOKEN ni o'rnating.")
        return
    
    application = Application.builder().token(TELEGRAM_BOT_TOKEN).build()
    
    application.add_handler(CommandHandler("start", start))
    application.add_handler(CommandHandler("admin", admin_panel))
    application.add_handler(CallbackQueryHandler(button_callback))
    application.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    application.add_handler(MessageHandler(filters.PHOTO, handle_photo))
    
    application.add_error_handler(error_handler)
    
    print("Bot ishga tushdi! Fayllarni qabul qilish uchun tayyor...")
    application.run_polling(allowed_updates=Update.ALL_TYPES, drop_pending_updates=True)

if __name__ == '__main__':
    main()
